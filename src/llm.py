import os
import requests
import json
import pandas as pd
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/{}:generateContent"
GEMINI_MODEL = "gemini-2.0-flash"

def prompt_gemini_cluster(text_data: str) -> dict:
    if not GEMINI_API_KEY:
        print("WARNING: GEMINI_API_KEY not found. Mocking newsletter.")
        return {"headline": "Mocked Headline", "bullet_points": ["Mocked bullet point"], "top_2_links": []}

    url = f"{GEMINI_URL.format(GEMINI_MODEL)}?key={GEMINI_API_KEY}"
    
    system_prompt = f"""
    You are an expert FMCG financial analyst. 
    You have been provided with a JSON array of news articles that all relate to the EXACT SAME real-world event or deal.
    
    CRITICAL INSTRUCTION:
    You MUST output a valid JSON object ONLY. No markdown, no conversational text.
    Your output must match this exact schema:
    {{
        "headline": "A single, highly professional business headline summarizing this unified event",
        "bullet_points": [
            "A concise bullet point capturing key financial facts or deal structure",
            "A concise bullet point capturing strategic impact, rationale, or locations"
        ],
        "sources": [
            {{"name": "Name of the publisher (e.g., Reuters, Bloomberg)", "url": "exact url"}}
        ]
    }}
    
    RULES:
    1. Synthesize all the provided articles into one coherent view of the event.
    2. Write NO MORE THAN 3-4 bullet points.
    3. The sources array must contain maximum 2 top sources from the provided list.
    
    Input JSON Records (Articles for this Event):
    {text_data}
    """
    
    payload = {
        "contents": [{"parts": [{"text": system_prompt}]}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "maxOutputTokens": 8192,
            "temperature": 0.2
        }
    }
    
    try:
        response = requests.post(url, json=payload)
        response.raise_for_status()
        data = response.json()
        raw_text = data['candidates'][0]['content']['parts'][0]['text'].strip()
        
        if raw_text.startswith("```"):
            raw_text = raw_text.split("```", 2)[-1]
            if raw_text.startswith("json"):
                raw_text = raw_text[4:]
            raw_text = raw_text.rsplit("```", 1)[0]
        raw_text = raw_text.strip()
        
        try:
            return json.loads(raw_text)
        except json.JSONDecodeError as je:
            import re
            match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except Exception:
                    pass
            raise je

    except Exception as e:
        print(f"Error parsing Gemini response: {e}")
        return {"headline": "Error Processing Cluster", "bullet_points": [str(e)], "sources": []}

def generate_newsletter(topics_json: list, df: pd.DataFrame, output_dir: str = "output"):
    os.makedirs(output_dir, exist_ok=True)
    df = df.copy()
    
    csv_path = os.path.join(output_dir, "advanced_nlp_clustered_topics.csv")
    df.to_csv(csv_path, index=False)
    print(f"Clustered NLP dataset exported to {csv_path}")

    safe_topics = []
    for t in topics_json:
        safe_t = dict(t)
        safe_t['topic_title'] = safe_t.get('topic_title', '').replace('\n', ' ').replace('\r', '')
        
        safe_articles = []
        for art in safe_t.get('articles', []):
            safe_art = dict(art)
            safe_art['title'] = safe_art.get('title', '').replace('\n', ' ')
            safe_art['summary'] = safe_art.get('summary', '').replace('\n', ' ').replace('"', "'")
            safe_articles.append(safe_art)
            
        safe_t['articles'] = safe_articles
        safe_topics.append(safe_t)
        
    # --- Generate Auxiliary Cluster Breakdown Report ---
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    breakdown_path = os.path.join(output_dir, "Cluster_Breakdown_Report.docx")
    b_doc = Document()
    b_header = b_doc.add_heading('FMCG Cluster Analytics & Source Breakdown', 0)
    b_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, topic in enumerate(safe_topics, 1):
        c_title = topic.get('topic_title', 'Unknown Topic')
        b_doc.add_heading(f"Cluster {i}: {c_title}", level=2)
        b_doc.add_paragraph(f"Total Articles in Cluster: {len(topic.get('articles', []))}\n")
        
        for art in topic.get('articles', []):
            p = b_doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{str(art.get('source', 'Unknown')).upper()}] ").bold = True
            p.add_run(str(art.get('title', 'No Title')))
            p_link = b_doc.add_paragraph(str(art.get('link', '')))
            p_link.style.font.size = Pt(8)
            
    b_doc.save(breakdown_path)
    print(f"Exported underlying raw cluster architecture to {breakdown_path}")
    # ---------------------------------------------------
        
    print(f"\n[Summary] Total valid deal clusters formed: {len(safe_topics)}")
    print("--- Top Clusters Breakdown ---")
    for i, topic in enumerate(safe_topics[:10], 1):
        print(f"  {i}. {topic.get('topic_title', 'Unknown Topic')[:60]}... -> {len(topic.get('articles', []))} articles")
    print("------------------------------\n")
    
    print("Drafting advanced NLP newsletter via Gemini (Cluster-Based)...")
    all_deals = []
    
    # Process up to 10 clusters (1 API call per cluster)
    for topic in safe_topics[:10]:
        print(f"  -> Sending cluster '{topic['topic_title'][:50]}...' to Gemini...")
        json_data = json.dumps(topic['articles'], ensure_ascii=True)
        chunk_json = prompt_gemini_cluster(json_data)
        all_deals.append(chunk_json)
        
    print("  -> Generating final Executive Summary...")
    headlines = "\n".join([f"- {d.get('headline', '')}" for d in all_deals])
    exec_prompt = f"""
    You are an expert FMCG financial analyst.
    Write a 2-3 sentence executive summary synthesizing the overall landscape of the following {len(all_deals)} recent FMCG deal activities.
    Output strictly the raw text summary string. No json, no intro text.
    
    Headlines:
    {headlines}
    """
    
    exec_summary = "FMCG Deal activities mapped successfully alongside source analysis."
    if GEMINI_API_KEY:
        url = f"{GEMINI_URL.format(GEMINI_MODEL)}?key={GEMINI_API_KEY}"
        payload = {
            "contents": [{"parts": [{"text": exec_prompt}]}],
            "generationConfig": {"temperature": 0.2, "maxOutputTokens": 1000}
        }
        try:
            response = requests.post(url, json=payload)
            response.raise_for_status()
            exec_summary = response.json()['candidates'][0]['content']['parts'][0]['text'].strip()
        except Exception as e:
            print(f"Warning: Exec summary proxy failed: {e}")
            
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    docx_path = os.path.join(output_dir, "FMCG_Executive_Newsletter.docx")
    doc = Document()
    
    title = doc.add_heading('FMCG Deal Intelligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}\n")
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(exec_summary)
    
    doc.add_heading('Top Deals This Period', level=1)
    
    for deal in all_deals:
        doc.add_heading(deal.get('headline', 'Strategic FMCG Deal'), level=2)
        
        for bullet in deal.get('bullet_points', []):
            doc.add_paragraph(str(bullet), style='List Bullet')
            
        sources = deal.get('sources', [])
        if sources:
            p_line = doc.add_paragraph()
            run = p_line.add_run("\nSources:")
            run.bold = True
            for src in sources:
                name = src.get('name', 'Link')
                url = src.get('url', '')
                if url and str(url).startswith('http'):
                    doc.add_paragraph(f"{name}: {url}")
            
        doc.add_paragraph()
        
    doc.save(docx_path)
    print(f"Successfully compiled Business Newsletter to {docx_path}")
